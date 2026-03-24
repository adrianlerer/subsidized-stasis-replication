#!/usr/bin/env python3
"""
Generate a professionally formatted academic .docx from the markdown draft
of "Subsidized Stasis" (Lerer, 2026).

Usage:
    python generate_docx.py
"""

import re
import os
import copy
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── paths ────────────────────────────────────────────────────────────────────
MD_PATH = Path("/Users/adria1/legal-evolution-unified/research/federalismo-fiscal-egt/paper_draft.md")
OUT_PRIMARY = Path("/Users/adria1/legal-evolution-unified/research/federalismo-fiscal-egt/Lerer_2026_Subsidized_Stasis.docx")
OUT_COPY = Path("/Users/adria1/Downloads/subsidized-stasis-replication/paper/Lerer_2026_Subsidized_Stasis.docx")

FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(12)
RUNNING_HEADER = "Subsidized Stasis \u2014 Lerer 2026"


# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_border(cell, **kwargs):
    """Set cell borders. Usage: set_cell_border(cell, top={"sz": 6, "val": "single"}, ...)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", 4)}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_run_font(run, size=BODY_SIZE, bold=False, italic=False):
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_formatted_text(paragraph, text, size=BODY_SIZE, bold=False, italic=False):
    """Add a run with specified formatting."""
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return run


def parse_inline_formatting(paragraph, text, size=BODY_SIZE, base_bold=False, base_italic=False):
    """Parse markdown bold/italic inline and add runs to a paragraph."""
    # Pattern matches **bold**, *italic*, ***bold-italic***
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)')
    pos = 0
    for m in pattern.finditer(text):
        # Add text before match
        if m.start() > pos:
            add_formatted_text(paragraph, text[pos:m.start()], size=size,
                               bold=base_bold, italic=base_italic)
        if m.group(2):  # ***bold-italic***
            add_formatted_text(paragraph, m.group(2), size=size, bold=True, italic=True)
        elif m.group(3):  # **bold**
            add_formatted_text(paragraph, m.group(3), size=size, bold=True, italic=base_italic)
        elif m.group(4):  # *italic*
            add_formatted_text(paragraph, m.group(4), size=size, bold=base_bold, italic=True)
        pos = m.end()
    # Remaining text
    if pos < len(text):
        add_formatted_text(paragraph, text[pos:], size=size, bold=base_bold, italic=base_italic)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=None, double=True):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    elif double:
        pf.line_spacing = Pt(24)  # double-spaced for 12pt


def add_page_number(doc):
    """Add page number at bottom center of each section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # PAGE field
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)
        for r in [run, run2, run3]:
            r.font.name = FONT_NAME
            r.font.size = Pt(10)


def add_running_header(doc, text):
    """Add running header to all sections."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.font.italic = True


# ── markdown parser ──────────────────────────────────────────────────────────

def parse_markdown(md_text):
    """Parse markdown into a list of blocks."""
    lines = md_text.split("\n")
    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Horizontal rule
        if re.match(r'^---+\s*$', line):
            i += 1
            continue

        # Empty line
        if line.strip() == "":
            i += 1
            continue

        # Table detection
        if "|" in line and i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1]):
            table_lines = [line]
            i += 1
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            blocks.append({"type": "table", "lines": table_lines})
            continue

        # Headers
        h_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            blocks.append({"type": "header", "level": level, "text": text})
            i += 1
            continue

        # Code/formula blocks (indented with 4 spaces)
        if line.startswith("    ") and not line.strip().startswith("-") and not line.strip().startswith("*"):
            code_lines = []
            while i < len(lines) and (lines[i].startswith("    ") or lines[i].strip() == ""):
                if lines[i].strip() == "" and (i + 1 >= len(lines) or not lines[i + 1].startswith("    ")):
                    break
                code_lines.append(lines[i])
                i += 1
            blocks.append({"type": "formula", "text": "\n".join(l.strip() for l in code_lines if l.strip())})
            continue

        # Numbered list items
        list_match = re.match(r'^(\d+)\.\s+(.*)', line)
        if list_match:
            items = []
            while i < len(lines):
                lm = re.match(r'^(\d+)\.\s+(.*)', lines[i])
                if lm:
                    item_text = lm.group(2)
                    i += 1
                    # Continuation lines
                    while i < len(lines) and lines[i].strip() and not re.match(r'^(\d+)\.\s+', lines[i]) and not re.match(r'^#{1,4}\s+', lines[i]) and not lines[i].startswith("    ") and "|" not in lines[i]:
                        item_text += " " + lines[i].strip()
                        i += 1
                    items.append(item_text)
                else:
                    break
            blocks.append({"type": "numbered_list", "items": items})
            continue

        # Bullet list items
        bullet_match = re.match(r'^[-*]\s+(.*)', line)
        if bullet_match:
            items = []
            while i < len(lines):
                bm = re.match(r'^[-*]\s+(.*)', lines[i])
                if bm:
                    item_text = bm.group(1)
                    i += 1
                    while i < len(lines) and lines[i].strip() and not re.match(r'^[-*]\s+', lines[i]) and not re.match(r'^#{1,4}\s+', lines[i]):
                        item_text += " " + lines[i].strip()
                        i += 1
                    items.append(item_text)
                else:
                    break
            blocks.append({"type": "bullet_list", "items": items})
            continue

        # Paragraph - accumulate consecutive non-empty lines
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r'^#{1,4}\s+', lines[i]) and not re.match(r'^---+\s*$', lines[i]) and "|" not in lines[i] and not re.match(r'^\d+\.\s+', lines[i]) and not re.match(r'^[-*]\s+', lines[i]) and not lines[i].startswith("    "):
            para_lines.append(lines[i])
            i += 1
        blocks.append({"type": "paragraph", "text": " ".join(l.strip() for l in para_lines)})

    return blocks


def parse_table(table_lines):
    """Parse markdown table lines into header row and data rows."""
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    # rows[0] = header, rows[1] = separator, rows[2:] = data
    if len(rows) < 2:
        return [], []
    header = rows[0]
    separator_idx = None
    for idx, row in enumerate(rows):
        if all(re.match(r'^[-:]+$', c) for c in row):
            separator_idx = idx
            break
    if separator_idx is not None:
        data = rows[separator_idx + 1:]
    else:
        data = rows[1:]
    return header, data


# ── document builder ─────────────────────────────────────────────────────────

def build_docx(blocks):
    doc = Document()

    # ── Page setup ───────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.paragraph_format.line_spacing = Pt(24)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    rPr = style.element.rPr
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(rPr)

    # ── Title page ───────────────────────────────────────────────────────
    # Extract metadata from blocks
    title_text = ""
    subtitle_text = ""
    author_info = {}
    abstract_text = ""
    keywords_text = ""
    jel_text = ""
    online_text = ""

    # Find title, subtitle, author info, abstract
    body_start = 0
    in_abstract = False
    abstract_parts = []

    for idx, block in enumerate(blocks):
        if block["type"] == "header" and block["level"] == 1:
            title_text = block["text"]
            continue
        if block["type"] == "header" and block["level"] == 2 and block["text"].startswith("Why "):
            subtitle_text = block["text"]
            continue
        if block["type"] == "paragraph":
            text = block["text"]
            if text.startswith("**Author:**"):
                author_info["author"] = text.replace("**Author:**", "").strip()
                continue
            if text.startswith("**Affiliation:**"):
                author_info["affiliation"] = text.replace("**Affiliation:**", "").strip()
                continue
            if text.startswith("**Email:**"):
                author_info["email"] = text.replace("**Email:**", "").strip()
                continue
            if text.startswith("**ORCID:**"):
                author_info["orcid"] = text.replace("**ORCID:**", "").strip()
                continue
            if text.startswith("**Date:**"):
                author_info["date"] = text.replace("**Date:**", "").strip()
                continue
            if text.startswith("**Keywords:**"):
                keywords_text = text.replace("**Keywords:**", "").strip()
                continue
            if text.startswith("**JEL Codes:**"):
                jel_text = text.replace("**JEL Codes:**", "").strip()
                continue
            if text.startswith("**Online Materials:**"):
                online_text = text.replace("**Online Materials:**", "").strip()
                continue

        if block["type"] == "header" and block["level"] == 2 and block["text"] == "ABSTRACT":
            in_abstract = True
            continue
        if in_abstract:
            if block["type"] == "header":
                in_abstract = False
                body_start = idx
                break
            if block["type"] == "paragraph":
                abstract_parts.append(block["text"])
                continue
        # Find where main body starts (Section 1)
        if block["type"] == "header" and block["level"] == 2 and "INTRODUCTION" in block["text"].upper():
            body_start = idx
            break

    abstract_text = "\n\n".join(abstract_parts)

    # ── Build title page ─────────────────────────────────────────────────
    # Add several empty paragraphs for spacing
    for _ in range(4):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, double=False, line_spacing=Pt(12))

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=12, double=False, line_spacing=Pt(28))
    add_formatted_text(p, title_text, size=Pt(16), bold=True)

    # Subtitle
    if subtitle_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=24, double=False, line_spacing=Pt(20))
        add_formatted_text(p, subtitle_text, size=Pt(13), italic=True)

    # Author
    for _ in range(2):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, double=False, line_spacing=Pt(12))

    if "author" in author_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=6, double=False, line_spacing=Pt(16))
        add_formatted_text(p, author_info["author"], size=Pt(14), bold=True)

    if "affiliation" in author_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=4, double=False, line_spacing=Pt(14))
        add_formatted_text(p, author_info["affiliation"], size=Pt(12))

    if "email" in author_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=4, double=False, line_spacing=Pt(14))
        add_formatted_text(p, author_info["email"], size=Pt(11))

    if "orcid" in author_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=4, double=False, line_spacing=Pt(14))
        add_formatted_text(p, "ORCID: " + author_info["orcid"], size=Pt(11))

    if "date" in author_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=4, double=False, line_spacing=Pt(14))
        add_formatted_text(p, author_info["date"], size=Pt(12))

    # Page break after title page
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

    # ── Abstract page ────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=12, double=False, line_spacing=Pt(14))
    add_formatted_text(p, "ABSTRACT", size=Pt(14), bold=True)

    # Abstract body - single spaced
    for abs_para in abstract_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(p, after=6, double=False, line_spacing=Pt(14))
        p.paragraph_format.first_line_indent = Inches(0.5)
        parse_inline_formatting(p, abs_para, size=Pt(11))

    # Keywords
    if keywords_text:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=12, after=4, double=False, line_spacing=Pt(14))
        add_formatted_text(p, "Keywords: ", size=Pt(11), bold=True)
        parse_inline_formatting(p, keywords_text, size=Pt(11))

    # JEL
    if jel_text:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=4, double=False, line_spacing=Pt(14))
        add_formatted_text(p, "JEL Codes: ", size=Pt(11), bold=True)
        parse_inline_formatting(p, jel_text, size=Pt(11))

    # Online materials
    if online_text:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=4, double=False, line_spacing=Pt(14))
        add_formatted_text(p, "Online Materials: ", size=Pt(11), bold=True)
        add_formatted_text(p, online_text, size=Pt(11))

    # Page break after abstract
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

    # ── Main body ────────────────────────────────────────────────────────
    section_counter = [0, 0, 0]  # for numbered sections
    table_num = 0

    for idx in range(body_start, len(blocks)):
        block = blocks[idx]

        if block["type"] == "header":
            level = block["level"]
            text = block["text"]

            # Check if text already has section number
            has_number = re.match(r'^(\d+\.?\d*\.?\d*)\s+', text)

            if level == 2:  # Main sections (## X. TITLE)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_spacing(p, before=24, after=12)
                add_formatted_text(p, text.upper() if text == text.upper() or "REFERENCES" in text.upper() else text, size=Pt(14), bold=True)

            elif level == 3:  # Subsections (### X.Y Title)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_spacing(p, before=18, after=8)
                add_formatted_text(p, text, size=Pt(13), bold=True)

            elif level == 4:  # Sub-subsections
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_spacing(p, before=12, after=6)
                add_formatted_text(p, text, size=Pt(12), bold=True, italic=True)

        elif block["type"] == "paragraph":
            text = block["text"]
            # Skip metadata paragraphs that were already used
            if any(text.startswith(prefix) for prefix in
                   ["**Author:**", "**Affiliation:**", "**Email:**",
                    "**ORCID:**", "**Date:**", "**Keywords:**",
                    "**JEL Codes:**", "**Online Materials:**"]):
                continue

            # Source lines under tables
            if text.startswith("*Sources:*") or text.startswith("*Source:*"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_spacing(p, after=6, double=False, line_spacing=Pt(14))
                parse_inline_formatting(p, text, size=Pt(10))
                continue

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, after=0)
            p.paragraph_format.first_line_indent = Inches(0.5)
            parse_inline_formatting(p, text)

        elif block["type"] == "formula":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=6, after=6, double=False, line_spacing=Pt(18))
            add_formatted_text(p, block["text"], size=Pt(11), italic=True)

        elif block["type"] == "numbered_list":
            for i, item in enumerate(block["items"], 1):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                set_paragraph_spacing(p, after=4)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.3)
                add_formatted_text(p, f"{i}. ", bold=True)
                parse_inline_formatting(p, item)

        elif block["type"] == "bullet_list":
            for item in block["items"]:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                set_paragraph_spacing(p, after=4)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                add_formatted_text(p, "\u2022 ")
                parse_inline_formatting(p, item)

        elif block["type"] == "table":
            header, data = parse_table(block["lines"])
            if not header:
                continue

            table_num += 1
            num_cols = len(header)
            num_rows = len(data) + 1  # +1 for header

            # Table caption
            # (tables in the paper have contextual labels from preceding text)

            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True

            # Set table style with borders
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
            borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                f'</w:tblBorders>'
            )
            tblPr.append(borders)

            # Header row
            for j, cell_text in enumerate(header):
                cell = table.cell(0, j)
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(p, double=False, line_spacing=Pt(14))
                parse_inline_formatting(p, cell_text, size=Pt(10), base_bold=True)
                # Header shading
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)

            # Data rows
            for r_idx, row_data in enumerate(data):
                for j, cell_text in enumerate(row_data):
                    if j >= num_cols:
                        break
                    cell = table.cell(r_idx + 1, j)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_paragraph_spacing(p, double=False, line_spacing=Pt(14))
                    parse_inline_formatting(p, cell_text, size=Pt(10))

            # Add spacing after table
            p = doc.add_paragraph()
            set_paragraph_spacing(p, after=6, double=False, line_spacing=Pt(6))

    # ── Headers and footers ──────────────────────────────────────────────
    add_running_header(doc, RUNNING_HEADER)
    add_page_number(doc)

    return doc


# ── main ─────────────────────────────────────────────────────────────────────

def main():
    print(f"Reading markdown from {MD_PATH}...")
    md_text = MD_PATH.read_text(encoding="utf-8")

    print("Parsing markdown...")
    blocks = parse_markdown(md_text)
    print(f"  Found {len(blocks)} blocks")

    print("Building DOCX...")
    doc = build_docx(blocks)

    # Save primary
    OUT_PRIMARY.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PRIMARY))
    print(f"  Saved: {OUT_PRIMARY}")

    # Save copy
    OUT_COPY.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_COPY))
    print(f"  Saved: {OUT_COPY}")

    print("Done!")


if __name__ == "__main__":
    main()
